from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from datetime import date

OUT = r"C:\Users\nGle-simoony2\Documents\ChatGPT\ai 자동화\AI_TC_자동수행_서비스_기획서.docx"
NAVY = "16324F"; BLUE = "2E74B5"; SKY = "EAF2F8"; PALE = "F4F6F9"
GRAY = "5E6B75"; LIGHT = "D9E2EA"; GREEN = "2E7D32"; AMBER = "9A6700"; RED = "B42318"

def set_font(run, size=None, bold=None, color=None, name="Malgun Gothic"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc.get_or_add_tcPr(); node = tc.first_child_found_in("w:tcMar")
    if node is None: node = OxmlElement("w:tcMar"); tc.append(node)
    for k,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        x=OxmlElement("w:"+k); x.set(qn("w:w"),str(v)); x.set(qn("w:type"),"dxa"); node.append(x)

def repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement("w:tblHeader"); e.set(qn("w:val"),"true"); trPr.append(e)

def fixed_table(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    tblPr=table._tbl.tblPr; tblW=tblPr.first_child_found_in("w:tblW")
    if tblW is None: tblW=OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"),"dxa")
    grid=table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for w in widths:
        g=OxmlElement("w:gridCol"); g.set(qn("w:w"),str(w)); grid.append(g)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            cell.width=Inches(widths[i]/1440); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
            tcW=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"),str(widths[i])); tcW.set(qn("w:type"),"dxa")

def table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i], "E8EEF5"); p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); set_font(r,9,bold=True,color=NAVY)
    repeat_header(t.rows[0])
    for vals in rows:
        cs=t.add_row().cells
        for i,val in enumerate(vals):
            p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0)
            r=p.add_run(str(val)); set_font(r,8.7,color="27323A")
    fixed_table(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f"Heading {level}"); p.add_run(text); return p

def para(doc, text, bold_prefix=None):
    p=doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); set_font(r,bold=True); r=p.add_run(text[len(bold_prefix):]); set_font(r)
    else: set_font(p.add_run(text))
    return p

def bullet(doc, text, level=0):
    p=doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2"); set_font(p.add_run(text)); return p

def num(doc, text):
    p=doc.add_paragraph(style="List Number"); set_font(p.add_run(text)); return p

def callout(doc, label, text, color=BLUE):
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; fixed_table(t,[9360]); shade(t.cell(0,0), PALE)
    p=t.cell(0,0).paragraphs[0]; r=p.add_run(label+"  "); set_font(r,10,bold=True,color=color); set_font(p.add_run(text),9.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

# Styles: standard_business_brief preset, Korean font override.
normal=doc.styles["Normal"]; normal.font.name="Malgun Gothic"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Malgun Gothic"); normal.font.size=Pt(10.5)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for nm,size,color,before,after in [("Title",25,NAVY,0,6),("Subtitle",13,GRAY,0,14),("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",11.5,NAVY,8,4)]:
    s=doc.styles[nm]; s.font.name="Malgun Gothic"; s._element.rPr.rFonts.set(qn("w:eastAsia"),"Malgun Gothic"); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=(nm.startswith("Heading")); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for nm in ["List Bullet","List Bullet 2","List Number"]:
    s=doc.styles[nm]; s.font.name="Malgun Gothic"; s._element.rPr.rFonts.set(qn("w:eastAsia"),"Malgun Gothic"); s.font.size=Pt(10); s.paragraph_format.space_after=Pt(4)

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_font(hp.add_run("AI 기반 TC 자동 수행 서비스 | 기획서"),8.5,color=GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
field=OxmlElement("w:fldSimple"); field.set(qn("w:instr"),"PAGE"); fp._p.append(field)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(70); p.paragraph_format.space_after=Pt(8)
set_font(p.add_run("PRODUCT CONCEPT & MVP PLAN"),10,bold=True,color=BLUE)
p=doc.add_paragraph(style="Title"); set_font(p.add_run("AI 기반 TC 자동 수행 서비스"),25,bold=True,color=NAVY)
p=doc.add_paragraph(style="Subtitle"); set_font(p.add_run("비정형 테스트 케이스를 해석해 화면을 탐색·조작·검증하는 다중 사용자 웹 플랫폼"),13,color=GRAY)
table(doc,["문서 구분","대상","작성 기준일"],[["서비스/기술 기획서","QA·개발·운영 조직",date.today().isoformat()]],[2100,4200,3060])
callout(doc,"핵심 제안", "브라우저 자동화로 기술·운영 가능성을 먼저 검증한 뒤 Android 에뮬레이터와 실기기로 확장한다. AI는 계획과 화면 해석을 담당하고, 실행 엔진은 허용된 행동만 수행한다.")
heading(doc,"의사결정 요약",1)
bullet(doc,"1차 MVP: 웹 페이지 대상, 자연어 TC 입력 → 단계 추론 → UI 탐색 → 실행 → 증적/결과 보고")
bullet(doc,"핵심 차별점: 사전에 selector를 모두 작성하지 않아도 텍스트·접근성 트리·스크린샷을 조합해 대상을 탐색")
bullet(doc,"안전 원칙: 결제·삭제·계정 변경 등 위험 행동은 차단 또는 승인 후 실행")
bullet(doc,"모바일 확장: Android 에뮬레이터 + Appium/ADB 기반으로 시작하고, 게임 특성에 따라 비전 중심 제어를 추가")
doc.add_page_break()

heading(doc,"1. 기획 배경 및 문제 정의",1)
para(doc,"현재의 UI 자동화는 테스트 케이스마다 selector, 좌표, 대기 조건과 검증식을 개발자가 직접 구현해야 한다. TC 형식이 제각각이고 화면이 자주 바뀌는 모바일 게임 환경에서는 자동화 제작·유지 비용이 빠르게 커진다.")
heading(doc,"해결할 문제",2)
bullet(doc,"자연어로 작성된 비정형 TC를 사람이 다시 자동화 코드로 옮겨야 한다.")
bullet(doc,"동일 기능도 기기 해상도, 언어, 팝업, 이벤트 배너에 따라 UI 위치가 바뀐다.")
bullet(doc,"실패 시 ‘왜 못 찾았는지’와 실제 화면 증적이 분리되어 재현이 어렵다.")
bullet(doc,"다수 인원이 동시에 사용하면 실행 자원, 계정, 테스트 데이터가 충돌한다.")
heading(doc,"제품 가설",2)
callout(doc,"가설", "AI가 TC 의도를 구조화하고 현재 화면에서 다음 행동 후보를 찾은 뒤, 제한된 도구로 실행하고 결과를 재검증하면 비정형 TC의 상당 부분을 낮은 작성 비용으로 자동화할 수 있다.")

heading(doc,"2. 목표와 범위",1)
table(doc,["구분","내용"],[
 ["제품 목표","비정형 TC를 입력하면 실행 계획, 자동 수행, 단계별 증적, 최종 판정을 제공"],
 ["MVP 목표","일반 웹 서비스에서 핵심 흐름 성공률과 운영 비용을 측정"],
 ["확장 목표","Android 게임의 메뉴·상점·우편함·설정 등 정형 UI 흐름 자동화"],
 ["비목표","초기 단계에서 모든 게임·모든 애니메이션·전투 플레이를 범용 자동화"],
 ["성공 기준","선정 TC 세트에서 높은 완주율, 재현 가능한 실패 로그, 사람이 납득 가능한 판정"]
],[1600,7760])
heading(doc,"권장 초기 대상",2)
bullet(doc,"웹: 로그인, 검색, 필터, 등록/수정, 권한별 메뉴 노출, 폼 검증")
bullet(doc,"모바일 게임: 로비 진입, 우편함 보상 수령, 상점 탭 이동, 설정 변경, 특정 재화/문구 확인")
bullet(doc,"후순위: 실시간 전투, 연속 드래그·멀티터치, 프레임 단위 타이밍, 랜덤성이 큰 콘텐츠")

heading(doc,"3. 주요 사용자와 사용 시나리오",1)
table(doc,["사용자","필요","대표 시나리오"],[
 ["QA 담당자","코드 없이 반복 실행","TC 붙여넣기 → 대상 환경 선택 → 실행 → 실패 증적 검토"],
 ["QA 리드","품질·비용 통제","프로젝트별 성공률, 재시도율, AI 비용, 위험 행동 정책 관리"],
 ["개발자","재현 정보 확보","실패 단계의 화면, 콘솔/네트워크, AI 판단 근거와 이벤트 로그 확인"],
 ["관리자","다중 사용자 운영","사용자/권한, 실행 슬롯, 기기 풀, 비밀정보, 감사 로그 관리"]
],[1500,2700,5160])

heading(doc,"4. 서비스 흐름",1)
for s in [
 "사용자가 프로젝트, 테스트 환경, 계정/기기, TC 원문을 선택한다.",
 "TC 해석기가 전제조건·행동·예상결과·종료조건을 구조화하고 불명확한 항목을 표시한다.",
 "실행 오케스트레이터가 격리된 브라우저 또는 모바일 세션을 배정한다.",
 "관찰기가 DOM/접근성 트리/스크린샷/현재 URL을 수집한다.",
 "AI 에이전트가 다음 행동 후보와 신뢰도를 생성한다.",
 "정책 엔진이 허용 여부를 검사하고 실행기가 클릭·입력·스크롤·대기를 수행한다.",
 "검증기가 예상 상태를 확인한다. 실패하면 제한 횟수 내에서 재탐색하거나 종료한다.",
 "단계별 스크린샷, 선택 근거, 실행 로그, 최종 Pass/Fail/Blocked를 리포트로 저장한다."]:
    num(doc,s)
callout(doc,"중요", "AI가 판단한 좌표를 무조건 클릭하지 않는다. 후보 요소, 신뢰도, 화면 변화, 허용 행동을 함께 확인하고 불확실성이 높으면 중단 또는 사용자 승인을 요청한다.", AMBER)

heading(doc,"5. 기능 요구사항",1)
heading(doc,"5.1 TC 입력 및 해석",2)
bullet(doc,"붙여넣기, 파일 업로드(CSV/XLSX/DOCX/TXT), 추후 Jira·TestRail 등 연동")
bullet(doc,"전제조건, 단계, 입력값, 기대결과, 테스트 데이터, 위험 행동을 자동 추출")
bullet(doc,"원문과 AI 구조화 결과를 함께 보이고 사용자가 실행 전 수정 가능")
bullet(doc,"불명확한 표현은 가정으로 숨기지 않고 ‘확인 필요’로 표시")
heading(doc,"5.2 실행",2)
bullet(doc,"클릭, 텍스트 입력, 키 입력, 스크롤, 대기, 뒤로가기, 새로고침, 스크린샷")
bullet(doc,"텍스트/role/label/접근성 ID/시각적 유사도를 결합한 요소 탐색")
bullet(doc,"팝업 닫기, 로딩 대기, 동일 의미 버튼 재탐색 등 제한적 복구")
bullet(doc,"전체 실행, 한 단계 실행, 일시정지, 재개, 즉시 중단")
heading(doc,"5.3 판정 및 리포트",2)
bullet(doc,"Pass, Fail, Blocked, Needs Review 네 가지 상태")
bullet(doc,"각 단계의 전·후 화면, 실행 행동, 대상 요소, 소요시간, 신뢰도 기록")
bullet(doc,"실패 원인 분류: 요소 미발견, 환경 오류, 데이터 오류, assertion 실패, 정책 차단, AI 불확실")
bullet(doc,"HTML/PDF/JSON 내보내기와 이슈 트래커 첨부를 확장 기능으로 제공")
heading(doc,"5.4 다중 사용자",2)
bullet(doc,"조직/프로젝트/역할 기반 접근 제어, SSO 연동 확장")
bullet(doc,"실행 큐, 동시 실행 제한, 우선순위, 타임아웃, 사용자별 비용 한도")
bullet(doc,"브라우저·기기·테스트 계정 예약과 세션별 격리")
bullet(doc,"비밀정보는 Vault/KMS 계열 저장소에 보관하고 프롬프트·로그에서 마스킹")

heading(doc,"6. 화면 구성",1)
table(doc,["화면","핵심 구성"],[
 ["대시보드","프로젝트별 실행 현황, 성공률, 최근 실패, 큐·기기 상태"],
 ["TC 작성/가져오기","원문 편집기, 구조화 미리보기, 변수/데이터, 위험 단계 표시"],
 ["실행 설정","환경·브라우저/기기·계정·AI 모델·최대 단계·재시도 정책"],
 ["실행 모니터","실시간 화면, 현재 단계, 다음 행동, 신뢰도, 승인/중단 버튼"],
 ["결과 상세","타임라인, 전후 스크린샷, 로그, 판정, 실패 원인, 재실행"],
 ["관리자","사용자·권한, 실행 풀, 비밀정보, 모델/비용, 감사 정책"]
],[1900,7460])

heading(doc,"7. 권장 시스템 구조",1)
para(doc,"웹 UI와 API 서버가 작업을 등록하고, 큐 기반 오케스트레이터가 격리된 실행 워커를 배정한다. AI 계층은 TC 해석과 다음 행동 결정을 담당하며, 실제 입력은 정책 엔진을 통과한 명령만 실행된다.")
table(doc,["계층","역할","후보 기술"],[
 ["Frontend","TC 작성, 실행 모니터, 리포트","Next.js/React, WebSocket 또는 SSE"],
 ["Backend API","인증, 프로젝트, 실행·결과 API","FastAPI 또는 NestJS, PostgreSQL"],
 ["Queue/Orchestrator","실행 예약, 재시도, 동시성 제어","Redis Queue/Celery/BullMQ; 성장 시 Kubernetes"],
 ["Web Worker","브라우저 관찰·조작","Playwright 우선"],
 ["Mobile Worker","Android 관찰·조작","Appium + UiAutomator2 + ADB, 에뮬레이터/디바이스 팜"],
 ["AI Gateway","모델 호출, 프롬프트 버전, 캐시, 비용·추적","멀티모달 LLM API, 구조화 출력"],
 ["Storage","스크린샷·영상·로그·리포트","Object Storage + PostgreSQL"],
 ["Observability","실패 분석, 성능·비용","OpenTelemetry, 로그/메트릭 대시보드"]
],[1650,3100,4610])
heading(doc,"AI 판단 루프",2)
table(doc,["단계","입력","출력/제약"],[
 ["Plan","TC 원문, 환경 정보","구조화된 단계와 성공 조건(JSON Schema)"],
 ["Observe","DOM/접근성 트리, 스크린샷, 최근 행동","현재 상태와 후보 UI 요소"],
 ["Decide","현재 목표, 후보, 실행 이력","허용된 단일 행동 + 대상 + 신뢰도"],
 ["Act","정책 통과 명령","Playwright/Appium 동작"],
 ["Verify","행동 전후 상태, 기대결과","통과/재탐색/중단 및 근거"]
],[1200,3800,4360])

heading(doc,"8. 웹 MVP와 모바일 게임의 차이",1)
table(doc,["항목","웹 MVP","모바일 게임"],[
 ["인식 데이터","DOM·접근성 트리 + 화면","접근성 트리가 빈약할 수 있어 화면 비전 비중 증가"],
 ["대상 지정","role/name/text/selector 사용 가능","좌표·템플릿·OCR·object detection 조합"],
 ["실행 환경","컨테이너 브라우저 확장 용이","에뮬레이터/실기기 수와 GPU·ADB 연결 제약"],
 ["변동성","레이아웃·A/B 테스트","애니메이션, 팝업, 해상도, 렌더링 엔진, 랜덤 이벤트"],
 ["권장 순서","기술/운영 모델 검증","정형 메뉴부터 제한된 게임군으로 확장"]
],[1500,3500,4360])
callout(doc,"권장 결론", "모바일 게임을 최종 목표로 유지하되, 웹 MVP로 에이전트 루프·실행 격리·리포트·다중 사용자 운영을 먼저 검증한다. 이후 Android 단일 게임/단일 해상도/정형 메뉴로 범위를 좁혀 확장한다.")

heading(doc,"9. 안전·보안·운영 정책",1)
table(doc,["위험","통제 방안"],[
 ["오클릭/무한 루프","최대 단계·시간·재시도·동일 행동 반복 제한, 화면 변화 감지"],
 ["결제/삭제/계정 변경","금지 목록, 도메인·화면별 allowlist, 사람 승인 게이트"],
 ["계정·개인정보 노출","전용 테스트 계정, 비밀정보 분리, 로그/스크린샷 마스킹, 보존기간"],
 ["프롬프트 인젝션","화면 텍스트를 명령이 아닌 관찰 데이터로 처리, 도구 권한 최소화"],
 ["테넌트 간 데이터 혼합","조직별 논리 분리, 실행 세션·저장 경로 격리, 권한 감사"],
 ["AI 오판","근거·신뢰도 저장, deterministic rule 우선, 고위험 단계 인간 승인"]
],[2200,7160])

heading(doc,"10. 데이터 모델(핵심)",1)
table(doc,["엔터티","주요 필드"],[
 ["Organization/User/Role","조직, 사용자, 프로젝트 권한, 비용 한도"],
 ["Project/Environment","대상 URL·앱 빌드, 허용 도메인, 정책, 변수"],
 ["TestCase/Version","원문, 구조화 단계, 데이터 세트, 버전, 작성자"],
 ["Execution/StepRun","상태, 워커, 시작·종료, 행동, 대상, 신뢰도, 오류"],
 ["Artifact","스크린샷, 영상, trace, 콘솔/네트워크 로그, 보존 기한"],
 ["Device/Account Lease","기기·계정 예약, 점유자, 만료, 초기화 상태"]
],[2400,6960])

heading(doc,"11. 단계별 개발 로드맵",1)
table(doc,["단계","기간 예시","범위","완료 조건"],[
 ["PoC","3~4주","Playwright 단일 사용자, 20~30개 웹 TC, 화면+DOM 기반 행동","대표 TC 성공/실패 원인과 한계 측정"],
 ["MVP","6~8주","웹 UI, 인증, 큐, 다중 워커, 리포트, 정책 엔진","다수 사용자 운영·감사·재실행 가능"],
 ["Pilot","4~6주","실서비스 스테이징, 계정/데이터 관리, CI 연동","선정 팀의 반복 회귀 테스트에 사용"],
 ["Mobile Alpha","6~10주","Android 단일 게임·해상도, Appium/ADB, 비전 fallback","정형 메뉴 TC 세트 완주율 검증"],
 ["Scale","지속","기기 풀, 조직 분리, 모델 라우팅, 비용 최적화","SLA·용량·보안 기준 충족"]
],[1100,1200,4300,2760])
para(doc,"※ 기간은 3~5명 규모의 제품/개발/QA 혼합 팀을 가정한 추정치이며, 대상 시스템 접근성·TC 난이도·보안 요구에 따라 달라진다.")

heading(doc,"12. KPI 및 검증 방법",1)
table(doc,["지표","정의","초기 목표 예시"],[
 ["TC 완주율","사람 개입 없이 종료 상태까지 도달한 비율","웹 선정 세트 80% 이상"],
 ["정확 판정률","사람 판정과 AI 판정이 일치한 비율","90% 이상"],
 ["요소 탐색 성공률","의도한 UI 대상을 올바르게 선택한 비율","90% 이상"],
 ["평균 개입 횟수","TC 1건당 승인·수정·복구 횟수","1회 미만"],
 ["재현 가능 실패율","증적만으로 원인 분류 가능한 실패 비율","95% 이상"],
 ["실행 비용","TC 1건당 모델·인프라 비용","기준선 수립 후 단계별 절감"]
],[2100,4300,2960])
para(doc,"목표 수치는 사업 기준이 아니라 파일럿 출발점이다. TC 난이도별로 분리해 측정하고, 성공한 쉬운 TC만 남기는 선택 편향을 방지해야 한다.")

heading(doc,"13. 주요 리스크와 대응",1)
table(doc,["리스크","영향","대응"],[
 ["비정형 TC가 지나치게 모호함","잘못된 계획·판정","사전 구조화 화면, 필수 성공조건, 확인 필요 표시"],
 ["화면 비전 호출 비용·지연","대량 실행 비경제적","DOM/접근성 우선, 변화 영역 crop, 캐시, 경량 모델 라우팅"],
 ["게임 UI 접근성 부족","대상 식별 실패","OCR·템플릿·비전 모델, 화면별 anchor, 지원 범위 명시"],
 ["동시 실행 자원 부족","대기 증가","큐·예약·우선순위·quota, 자동 확장 가능한 웹 워커"],
 ["AI 결과 재현성 부족","디버깅 어려움","모델/프롬프트/관찰 snapshot 버전 고정, trace 저장"],
 ["약관·보안 이슈","서비스 사용 제한","테스트 환경·자사 앱 우선, 사내 정책 및 게임 약관 검토"]
],[2100,2800,4460])

heading(doc,"14. MVP 백로그",1)
table(doc,["우선순위","기능","수용 기준"],[
 ["P0","TC 원문 → 구조화 계획","JSON schema 검증, 사용자 수정/승인"],
 ["P0","격리 브라우저 실행","세션별 쿠키·스토리지 분리, 강제 종료 가능"],
 ["P0","화면 관찰과 요소 후보화","DOM/접근성/스크린샷을 단계 로그에 연결"],
 ["P0","안전 정책 엔진","금지 행동 차단, 위험 행동 승인, 반복 제한"],
 ["P0","단계별 결과/증적","전후 화면, 행동, 근거, 오류, 최종 판정"],
 ["P1","다중 사용자 큐/권한","동시성 제한, 프로젝트 격리, 실행 우선순위"],
 ["P1","재실행/실패 단계부터 실행","동일 데이터·환경 조건 복원"],
 ["P2","Android 실행 워커","단일 앱/기기 프로파일의 정형 TC 지원"]
],[1200,3400,4760])

heading(doc,"15. 구현 원칙 및 API 활용 방향",1)
bullet(doc,"AI API는 ‘판단 계층’으로 사용하고 브라우저/기기 제어 권한은 서버의 제한된 도구 계층에 둔다.")
bullet(doc,"모델 응답은 자유 텍스트가 아니라 구조화된 action schema로 검증한 뒤 실행한다.")
bullet(doc,"매 단계 전체 화면을 고비용 모델에 보내지 말고, DOM 기반 후보 축소 → 필요 시 스크린샷 비전 호출 순서로 라우팅한다.")
bullet(doc,"모델 사업자 교체가 가능하도록 AI Gateway를 두고 프롬프트·모델·비용·평가 세트를 버전 관리한다.")
bullet(doc,"규칙 기반 검증(정확한 URL, 텍스트, 숫자, DOM 상태)은 AI 판정보다 우선한다.")
heading(doc,"권장 PoC 기술 조합",2)
table(doc,["영역","권장안"],[
 ["UI/API","Next.js + FastAPI 또는 NestJS"],
 ["브라우저","Playwright worker, trace/video/screenshot 활성화"],
 ["AI","멀티모달 모델 + 구조화 출력 + tool calling"],
 ["데이터","PostgreSQL, Redis, S3 호환 Object Storage"],
 ["배포","Docker Compose로 시작, 동시성 증가 시 Kubernetes/managed queue 검토"],
 ["모바일","Appium + UiAutomator2 + ADB + Android Emulator"]
],[2400,6960])

heading(doc,"16. PoC 실행안",1)
for s in [
 "대상 웹 서비스 1개와 대표 TC 20~30개를 선정하고 난이도(쉬움/중간/어려움)를 고정한다.",
 "결제·삭제 없이 로그인, 검색, 입력, 상태 확인 중심으로 범위를 제한한다.",
 "TC 구조화 schema와 실행 action schema를 먼저 정의한다.",
 "DOM/접근성 우선 탐색과 스크린샷 fallback을 구현한다.",
 "모든 실행의 화면·trace·AI 입출력·비용을 저장한다.",
 "사람 기준 정답과 비교해 완주율, 판정 정확도, 지연, 비용을 측정한다.",
 "실패 상위 원인을 개선한 뒤 모바일 Alpha 진입 여부를 결정한다."]:
    num(doc,s)

heading(doc,"17. 의사결정 필요 항목",1)
bullet(doc,"최초 적용 대상: 사내 웹/스테이징 서비스와 대표 TC 세트")
bullet(doc,"사용 조직과 동시 실행 규모, 데이터 보존 기간")
bullet(doc,"AI API의 외부 전송 가능 데이터 범위와 사내 보안 검토")
bullet(doc,"모바일 목표 게임, Android 지원 버전, 에뮬레이터/실기기 여부")
bullet(doc,"결제·삭제·계정 변경 등 고위험 행동의 승인 정책")
callout(doc,"Go 제안", "4주 PoC를 승인하고 웹 TC 20~30개로 기술 지표를 수집한다. PoC 종료 시 완주율·정확도·비용·실패 유형을 기준으로 MVP 및 모바일 Alpha 투자를 결정한다.", GREEN)

doc.core_properties.title="AI 기반 TC 자동 수행 서비스 기획서"
doc.core_properties.subject="비정형 테스트 케이스 해석 및 UI 자동 수행 플랫폼"
doc.core_properties.author="Codex"
doc.save(OUT)
print(OUT)
